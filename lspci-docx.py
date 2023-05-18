from docx import Document
import re
import sys

#https://github.com/rdobson/python-hwinfo/blob/ba93a112dac6863396a053636ea87df027daa5de/hwinfo/pci/lspci.py#L46

COLOMUN_COUNT = 3

LABEL_REGEX = r'[\w+\ \.\,\:\+\&\-\/\[\]\(\)\#]+'
CODE_REGEX = r'[0-9a-fA-F]{4}'
BUSID_REGEX = r'[0-9A-Za-z]{2}:[0-9A-Fa-z]{2}\.[0-9A-Fa-f]{1}'
REV_REGEX = r'[0-9a-fA-F]{2}'
PROG_REGEX = r'[0-9a-fA-F]{2}'

class LspciNNMMParser():
    """Parser object for the output of lspci -nnmm"""

    #02:00.1 "Ethernet controller [0200]" "Broadcom Corporation [14e4]" "NetXtreme II BCM5716 Gigabit Ethernet [163b]" -r20 "Dell [1028]" "Device [02a3]"

    ITEM_REGEXS = [
        r'(?P<pci_device_bus_id>(' + BUSID_REGEX + r'))\ "(?P<pci_device_class_name>' + LABEL_REGEX + r')\ \[(?P<pci_device_class>' + CODE_REGEX + r')\]"' \
        + r'\ "(?P<pci_vendor_name>' + LABEL_REGEX + r')\ \[(?P<pci_vendor_id>' + CODE_REGEX + r')\]"\ "(?P<pci_device_name>' + LABEL_REGEX + r')\ \[(?P<pci_device_id>' + CODE_REGEX + r')\]"' \
        + r'((\ -r)(?P<pci_device_rev>' + REV_REGEX + r'))*' \
        + r'((\ -p)(?P<pci_device_prog>' + PROG_REGEX + r'))*' \
        + r'\ .*"((?P<pci_subvendor_name>' + LABEL_REGEX + r')\ \[(?P<pci_subvendor_id>' + CODE_REGEX + r')\])*"\ "((?P<pci_subdevice_name>' + LABEL_REGEX + r')\ \[(?P<pci_subdevice_id>' + CODE_REGEX + r')\])*'
        ,
    ]
    ITEM_SEPERATOR = "\n"

class Table():
    def __init__(self):
        self.__rev_list = [] #(Optional)
        self.__dev_list = []  # 00:00.0
        self.__class_name_list = []  # Host bridge
        self.__ven_list = []  # Intel Corporation
        self.__code_list = []  # Vendor: [8086]
        self.__prog_list = []  # ProgIf: 01 (Optional)
        self.__pos = 0  # iterator over class

    def __next__(self):
        if self.__pos >= len(self.__dev_list) or \
                self.__pos >= len(self.__class_name_list) or \
                self.__pos >= len(self.__ven_list) or \
                self.__pos >= len(self.__code_list) or \
                self.__pos >= len(self.__prog_list) or \
                self.__pos >= len(self.__rev_list):

                array_size = (
                    len(self.__dev_list), 
                   len(self.__class_name_list), 
                   len(self.__ven_list), 
                   len(self.__code_list), 
                   len(self.__prog_list), 
                   len(self.__rev_list),
                   )
                print(array_size)

                raise StopIteration()

        result = {
            "dev": self.__dev_list[self.__pos],
            "rev": self.__rev_list[self.__pos],
            "class": self.__class_name_list[self.__pos],
            "ven": self.__ven_list[self.__pos],
            "prog": self.__prog_list[self.__pos],
            "code": self.__code_list[self.__pos],
        }
        self.__pos += 1
        return result

    def __iter__(self):
        return self

    def set_rev(self, rev):
        if rev == None:
            rev="-"
        self.__rev_list.append(rev)

    def set_dev(self, dev):   
        self.__dev_list.append(dev)

    def set_class(self, _class_name):
        self.__class_name_list.append(_class_name)

    def set_ven_code(self, ven):
        self.__ven_list.append(ven)

    def set_code(self, code):
        self.__code_list.append(code)

    def set_class_name(self, _class_name):
        self.__class_name_list.append(_class_name)

    def set_prog(self, prog):     
        if prog == None:
            prog="-"   
        self.__prog_list.append(prog)
            
    def get_rev(self):
        return self.__rev_list

    def get_dev(self):
        return self.__dev_list

    def get_class_name(self):
        return self.__class_name_list

    def get_ven(self):
        return self.__ven_list

    def get_ven_code(self):
        return self.__code_list

    def get_prog(self):
        return self.__prog_list
    
def combine_dicts(recs):
    """Combine a list of recs, appending values to matching keys"""
    if not recs:
        return None

    if len(recs) == 1:
        return recs.pop()
    new_rec = {}

    for rec in recs:
        for k, v in rec.iteritems():
            if k in new_rec:
                new_rec[k] = "%s, %s" % (new_rec[k], v)
            else:
                new_rec[k] = v
    return new_rec

def parse_input(Table, file_path):
    parser = LspciNNMMParser()
    rec = {}
    
    try:
        with open(file_path, "r") as f:
            for line in f:
                for regex in parser.ITEM_REGEXS:
                    matches = [m.groupdict() for m in re.finditer(regex, line.strip())]
                    mdicts = combine_dicts(matches)
                    if mdicts:
                        rec = dict(list(rec.items()) + list(mdicts.items()))
                        Table.set_dev(rec.get("pci_device_bus_id"))
                        Table.set_class_name(rec.get("pci_device_class_name"))
                        Table.set_ven_code(rec.get("pci_vendor_name"))
                        Table.set_code(rec.get("pci_vendor_id"))
                        Table.set_rev(rec.get("pci_device_rev"))
                        Table.set_prog(rec.get("pci_device_prog"))
    except OSError as error:
        print(error)
        sys.exit(-1)

def set_table(Table):
    slot_col = "Сведения о местонахождении"
    dev_col = "Устройство"
    descr_col = "Описание"

    document = Document()

    table = document.add_table(rows=1, cols=COLOMUN_COUNT)
    table.style = "Table Grid"
    table.style.font.name = "Times New Roman"

    hdr_cells = table.rows[0].cells
    hdr_cells[0].paragraphs[0].add_run(slot_col).bold = True
    hdr_cells[1].paragraphs[0].add_run(dev_col).bold = True
    hdr_cells[2].paragraphs[0].add_run(descr_col).bold = True

    for row in Table:
        row_cells = table.add_row().cells
        row_cells[0].text = row.get("dev")
        row_cells[1].text = row.get("class")
        row_cells[2].text = "Производитель: " + row.get("ven") + \
            "\nКод вендора: " + row.get("code") + \
                "\nRev: " + row.get("rev") + \
                "\nProgIf: " + row.get("prog")

    if (len(sys.argv)) == 2:
        document.save("/opt/pci-dev.docx")
    else:
        document.save("pci-dev.docx")

def main():
    if (len(sys.argv)) == 2:
        file_path = sys.argv[1]
    else:
        file_path = "/pci-dev.text"
        print("No file spec. Using: " + file_path)

    table = Table()
    parse_input(table,file_path)
    set_table(table)
    
    #print(table.get_ven())
    #print(table.get_rev())
    #print(table.get_dev())
    #print (table.get_class_name())
    #print(table.get_ven_code())
    #print(table.get_prog())
    #print(len(table.get_rev()))
main()
